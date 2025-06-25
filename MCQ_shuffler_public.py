import docx
import random
import re


def copy_paragraph_preserve(src_para, dest_doc, new_text=None):
    """
    Copy a source paragraph (with its formatting and runs) into dest_doc.
    If new_text is provided, replace the paragraph text with new_text (applied in a single run).
    """
    new_para = dest_doc.add_paragraph()
    # Copy paragraph-level formatting
    new_para.style = src_para.style
    new_para.alignment = src_para.alignment
    pf_src = src_para.paragraph_format
    pf_dest = new_para.paragraph_format
    pf_dest.left_indent = pf_src.left_indent
    pf_dest.right_indent = pf_src.right_indent
    pf_dest.first_line_indent = pf_src.first_line_indent
    pf_dest.space_before = pf_src.space_before
    pf_dest.space_after = pf_src.space_after
    pf_dest.line_spacing = pf_src.line_spacing

    # Copy runs, optionally replacing text
    if new_text is not None:
        # Use formatting of the first run if available
        first_run = src_para.runs[0] if src_para.runs else None
        new_run = new_para.add_run(new_text)
        if first_run:
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            if first_run.font.name:
                new_run.font.name = first_run.font.name
            if first_run.font.size:
                new_run.font.size = first_run.font.size
            if first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb
    else:
        for run in src_para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
    return new_para


def shuffle_doc(input_path, output_path):
    # Load the existing document
    doc = docx.Document(input_path)

    # Parse blocks of question, answers, and separator
    blocks = []
    current = None
    for para in doc.paragraphs:
        raw = para.text
        text = raw.strip()
        if re.match(r'^\d+\.\s', text):
            # Start new question block
            if current:
                blocks.append(current)
            current = {'question': para, 'answers': [], 'sep': None}
        elif current and re.match(r'^[A-D]\.\s', text):
            # Answer choice
            current['answers'].append(para)
        elif current and raw == '':
            # Blank paragraph separator
            if current['sep'] is None:
                current['sep'] = para
        elif current:
            # Continuation line: append to last paragraph's runs
            target = current['answers'][-1] if current['answers'] else current['question']
            target.add_run('\n' + raw)
    if current:
        blocks.append(current)

    # Shuffle answers within each question
    for blk in blocks:
        random.shuffle(blk['answers'])

    # Shuffle the order of questions
    random.shuffle(blocks)

    # Create a new document preserving styles
    new_doc = docx.Document()

    # Write shuffled content with consistent numbering
    for idx, blk in enumerate(blocks, start=1):
        # Question with updated numbering
        q_text = re.sub(r'^\d+\.\s*', f'{idx}. ', blk['question'].text)
        copy_paragraph_preserve(blk['question'], new_doc, new_text=q_text)

        # Answers relabeled A-D
        for a_idx, ans_para in enumerate(blk['answers']):
            label = chr(ord('A') + a_idx)
            a_text = re.sub(r'^[A-D]\.\s*', f'{label}. ', ans_para.text)
            copy_paragraph_preserve(ans_para, new_doc, new_text=a_text)

        # Copy original blank separator or insert a blank line
        if blk['sep']:
            copy_paragraph_preserve(blk['sep'], new_doc)
        else:
            new_doc.add_paragraph()

    # Save the shuffled document
    new_doc.save(output_path)


if __name__ == '__main__':
    shuffle_doc('sample_input_questions.docx', 'shuffled_output_questions.docx')
